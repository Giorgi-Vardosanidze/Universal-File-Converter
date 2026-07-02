from pathlib import Path

from docx import Document
from pdf2docx import Converter
import PyPDF2
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.lib.styles import getSampleStyleSheet

from app.core.word_session import WordSession
from app.core.constants import WORD_PDF_FORMAT


class FileConverter:
    def __init__(self, log_callback=None):
        self.log_callback = log_callback

    def _word_export_pdf(self, src: Path, dst: Path):
        app = WordSession.get()
        doc = None
        try:
            doc = app.Documents.Open(str(src), ReadOnly=True)
            doc.ExportAsFixedFormat(OutputFileName=str(dst), ExportFormat=WORD_PDF_FORMAT)
        finally:
            if doc is not None:
                doc.Close(False)

    def doc_to_pdf(self, src, dst):
        self._word_export_pdf(src, dst)

    def docx_to_pdf(self, src, dst):
        self._word_export_pdf(src, dst)

    def doc_to_txt(self, src, dst):
        app = WordSession.get()
        doc = None
        try:
            doc = app.Documents.Open(str(src), ReadOnly=True)
            dst.write_text(doc.Content.Text, encoding="utf-8")
        finally:
            if doc is not None:
                doc.Close(False)

    def docx_to_txt(self, src, dst):
        doc = Document(str(src))
        text = "\n".join(p.text for p in doc.paragraphs)
        dst.write_text(text, encoding="utf-8")

    def txt_to_docx(self, src, dst):
        doc = Document()
        with open(src, "r", encoding="utf-8") as f:
            for line in f:
                doc.add_paragraph(line.rstrip("\n"))
        doc.save(str(dst))

    def txt_to_pdf(self, src, dst):
        styles = getSampleStyleSheet()
        story = []
        with open(src, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                story.append(Paragraph(line if line else "&nbsp;", styles["Normal"]))
        doc = SimpleDocTemplate(str(dst))
        doc.build(story)

    def pdf_to_txt(self, src, dst):
        reader = PyPDF2.PdfReader(str(src))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        dst.write_text(text, encoding="utf-8")

    def pdf_to_docx(self, src, dst):
        cv = Converter(str(src))
        try:
            cv.convert(str(dst))
        finally:
            cv.close()

    def pdf_to_doc(self, src, dst):
        temp_docx = dst.with_suffix(".tmp.docx")
        self.pdf_to_docx(src, temp_docx)
        app = WordSession.get()
        doc = None
        try:
            doc = app.Documents.Open(str(temp_docx))
            doc.SaveAs2(str(dst), FileFormat=0)
        finally:
            if doc is not None:
                doc.Close(False)
            if temp_docx.exists():
                temp_docx.unlink()

    def convert_file(self, src: Path, dst: Path, target_ext: str):
        s = src.suffix.lower().lstrip(".")
        t = target_ext

        if s == "doc" and t == "pdf":
            self.doc_to_pdf(src, dst)
        elif s == "docx" and t == "pdf":
            self.docx_to_pdf(src, dst)
        elif s == "docx" and t == "txt":
            self.docx_to_txt(src, dst)
        elif s == "doc" and t == "txt":
            self.doc_to_txt(src, dst)
        elif s == "txt" and t == "docx":
            self.txt_to_docx(src, dst)
        elif s == "txt" and t == "pdf":
            self.txt_to_pdf(src, dst)
        elif s == "pdf" and t == "txt":
            self.pdf_to_txt(src, dst)
        elif s == "pdf" and t == "docx":
            self.pdf_to_docx(src, dst)
        elif s == "pdf" and t == "doc":
            self.pdf_to_doc(src, dst)
        else:
            raise ValueError(f"მხარდაუჭერელი კონვერტაცია: .{s} -> .{t}")
